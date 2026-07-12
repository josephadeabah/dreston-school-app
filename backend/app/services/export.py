"""Generates branded, print-ready PDF and Word reports.

Kept deliberately simple: every export is "a title, a subtitle, and a
table" — that covers rosters, attendance sheets, payment lists, and
balance summaries without needing a bespoke layout per report.
"""

from datetime import datetime
from io import BytesIO

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import (
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

SCHOOL_NAME = "Dreston Elite Montessori School"
SCHOOL_MOTTO = "The fear of the Lord is the beginning of wisdom."

VIOLET = "#6B429F"
PLUM = "#3B1F45"
GOLD = "#C79A56"
BLUSH = "#FBEAF2"


def _generated_line() -> str:
    return f"Generated {datetime.now().strftime('%d %B %Y, %I:%M %p')}"


def build_pdf(
    title: str,
    subtitle: str | None,
    headers: list[str],
    rows: list[list[str]],
    column_widths: list[float] | None = None,
) -> bytes:
    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        topMargin=1.8 * cm,
        bottomMargin=1.8 * cm,
        leftMargin=1.5 * cm,
        rightMargin=1.5 * cm,
    )

    school_style = ParagraphStyle(
        "School",
        fontName="Helvetica-Bold",
        fontSize=15,
        textColor=colors.HexColor(VIOLET),
    )
    motto_style = ParagraphStyle(
        "Motto",
        fontName="Helvetica-Oblique",
        fontSize=9,
        textColor=colors.HexColor(GOLD),
    )
    title_style = ParagraphStyle(
        "Title",
        fontName="Helvetica-Bold",
        fontSize=13,
        textColor=colors.HexColor(PLUM),
        spaceBefore=14,
    )
    subtitle_style = ParagraphStyle(
        "Subtitle", fontName="Helvetica", fontSize=10, textColor=colors.HexColor(PLUM)
    )
    meta_style = ParagraphStyle(
        "Meta", fontName="Helvetica", fontSize=8, textColor=colors.grey, spaceBefore=2
    )

    elements = [
        Paragraph(SCHOOL_NAME, school_style),
        Paragraph(SCHOOL_MOTTO, motto_style),
        Paragraph(title, title_style),
    ]
    if subtitle:
        elements.append(Paragraph(subtitle, subtitle_style))
    elements.append(Paragraph(_generated_line(), meta_style))
    elements.append(Spacer(1, 12))

    table_data = [headers] + rows if rows else [headers]
    table = Table(table_data, colWidths=column_widths, repeatRows=1)
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(VIOLET)),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                (
                    "ROWBACKGROUNDS",
                    (0, 1),
                    (-1, -1),
                    [colors.white, colors.HexColor(BLUSH)],
                ),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#E7A0C5")),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                ("LEFTPADDING", (0, 0), (-1, -1), 8),
            ]
        )
    )
    elements.append(table)

    if not rows:
        elements.append(Spacer(1, 10))
        elements.append(Paragraph("No records for this selection.", subtitle_style))

    doc.build(elements)
    return buffer.getvalue()


def build_docx(
    title: str,
    subtitle: str | None,
    headers: list[str],
    rows: list[list[str]],
) -> bytes:
    doc = Document()

    section = doc.sections[0]
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    school_p = doc.add_paragraph()
    school_run = school_p.add_run(SCHOOL_NAME)
    school_run.bold = True
    school_run.font.size = Pt(18)
    school_run.font.color.rgb = RGBColor(0x6B, 0x42, 0x9F)

    motto_p = doc.add_paragraph()
    motto_run = motto_p.add_run(SCHOOL_MOTTO)
    motto_run.italic = True
    motto_run.font.size = Pt(10)
    motto_run.font.color.rgb = RGBColor(0xC7, 0x9A, 0x56)

    title_p = doc.add_paragraph()
    title_run = title_p.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(14)
    title_run.font.color.rgb = RGBColor(0x3B, 0x1F, 0x45)
    title_p.space_before = Pt(14)

    if subtitle:
        sub_p = doc.add_paragraph()
        sub_run = sub_p.add_run(subtitle)
        sub_run.font.size = Pt(10.5)

    meta_p = doc.add_paragraph()
    meta_run = meta_p.add_run(_generated_line())
    meta_run.font.size = Pt(8.5)
    meta_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    doc.add_paragraph()

    if rows:
        table = doc.add_table(rows=1, cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Light Grid Accent 4"

        header_cells = table.rows[0].cells
        for i, h in enumerate(headers):
            header_cells[i].text = h
            for p in header_cells[i].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for r in p.runs:
                    r.bold = True

        for row in rows:
            cells = table.add_row().cells
            for i, val in enumerate(row):
                cells[i].text = str(val)
    else:
        doc.add_paragraph("No records for this selection.")

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
